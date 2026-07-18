#!/usr/bin/env python3
# Klon podania e-Curia Z PETROVHO WORDU: zachova formatovanie (Calibri 12, okraje, styly,
# tucne useky, textboxy s podpisom), vymeni len text za preklad + vlozi banner a cisty podpis.
# Usage: python3 build_docx_clone.py <LANG>
import sys, copy, shutil, os, subprocess, zipfile
from pypdf import PdfReader
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL

SRC = "/home/claude/podanie03/03_ANTWORT AUF DIE ABLEHNUNG DES ANTRAGS AUF ZUGANG ZU e-CURIA_SK.docx"
HERE = os.path.dirname(os.path.abspath(__file__))

# ---- preklady: index odseku -> [(text, bold), ...] ; nacitane z modulu jazyka
def load_lang(code):
    mod = {}
    exec(open(os.path.join(HERE, "lang_%s.py" % code.lower()), encoding="utf-8").read(), mod)
    return mod["PARA"], mod["LABELS"], mod["BOXES"]

WPNS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

def _float_right(run):
    """Premeni inline obrazok behu na PLAVAJUCI ukotveny vpravo hore (wrapNone)."""
    drawing = run._element.find(qn('w:drawing'))
    if drawing is None:
        return
    inline = drawing.find('{%s}inline' % WPNS)
    if inline is None:
        return
    anchor = OxmlElement('wp:anchor')
    for k, v in (("distT","0"),("distB","0"),("distL","114300"),("distR","114300"),
                 ("simplePos","0"),("relativeHeight","251658240"),("behindDoc","0"),
                 ("locked","0"),("layoutInCell","1"),("allowOverlap","1")):
        anchor.set(k, v)
    sp = OxmlElement('wp:simplePos'); sp.set("x","0"); sp.set("y","0"); anchor.append(sp)
    ph = OxmlElement('wp:positionH'); ph.set("relativeFrom","margin")
    al = OxmlElement('wp:align'); al.text = "right"; ph.append(al); anchor.append(ph)
    pv = OxmlElement('wp:positionV'); pv.set("relativeFrom","paragraph")
    off = OxmlElement('wp:posOffset'); off.text = "0"; pv.append(off); anchor.append(pv)
    for tag in ('{%s}extent' % WPNS, '{%s}effectExtent' % WPNS):
        el = inline.find(tag)
        if el is not None:
            anchor.append(copy.deepcopy(el))
    anchor.append(OxmlElement('wp:wrapNone'))
    for tag in ('{%s}docPr' % WPNS, '{%s}cNvGraphicFramePr' % WPNS):
        el = inline.find(tag)
        if el is not None:
            anchor.append(copy.deepcopy(el))
    gr = inline.find('{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
    if gr is not None:
        anchor.append(copy.deepcopy(gr))
    drawing.remove(inline)
    drawing.append(anchor)

def set_para(p_el, segs, base_rpr):
    """Prepise LEN textove behy odseku; ukotvenu grafiku (textboxy, podpis) necha na mieste."""
    def is_gfx(el):
        for ch in el.iter():
            ln = ch.tag.split('}')[-1]
            if ln in ("drawing", "pict", "AlternateContent", "object"):
                return True
        return False
    text_runs = [el for el in list(p_el)
                 if el.tag == qn('w:r') and el.find(qn('w:t')) is not None and not is_gfx(el)]
    idx = list(p_el).index(text_runs[0]) if text_runs else len(list(p_el))
    for r in text_runs:
        p_el.remove(r)
    new = []
    for text, bold in segs:
        r = OxmlElement('w:r')
        rpr = copy.deepcopy(base_rpr) if base_rpr is not None else OxmlElement('w:rPr')
        for tag in ('w:b', 'w:bCs'):
            el = rpr.find(qn(tag))
            if bold and el is None:
                rpr.append(OxmlElement(tag))
            elif not bold and el is not None:
                rpr.remove(el)
        r.append(rpr)
        t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
        r.append(t)
        new.append(r)
    for k, r in enumerate(new):
        p_el.insert(idx + k, r)

def first_rpr(p):
    for r in p.runs:
        if r._element.find(qn('w:rPr')) is not None:
            return r._element.find(qn('w:rPr'))
    return None

def build(code):
    PARA, LABELS, BOXES = load_lang(code)
    out_docx = os.path.join(HERE, "03_e-Curia_%s.docx" % code)
    shutil.copy(SRC, out_docx)

    # 1) cisty podpis (vysoke rozlisenie) namiesto defektneho v Worde
    tmp = out_docx + ".zip"
    zin = zipfile.ZipFile(out_docx, 'r')
    zout = zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED)
    clean = open(os.path.join(HERE, "sig_hires_clean.png"), 'rb').read()
    for it in zin.infolist():
        data = zin.read(it.filename)
        if it.filename == "word/media/image1.png":
            data = clean
        zout.writestr(it, data)
    zin.close(); zout.close(); shutil.move(tmp, out_docx)

    d = Document(out_docx)

    # 2) telo — vymena textu odsekov
    for i, p in enumerate(d.paragraphs):
        if i in PARA:
            set_para(p._element, PARA[i], first_rpr(p))
        elif i in LABELS:                       # odseky s hypertextovym e-mailom: len popisky
            for r in p.runs:
                for old, new in LABELS[i]:
                    if r.text.strip() == old.strip():
                        r.text = r.text.replace(old.strip(), new)

    # 3) PODPISOVY BLOK — v origináli su to PLAVAJUCE textboxy s pevnou poziciou; pri dlhsom
    # texte (FR/ES/IT) ich vytlaci pod spodny okraj a orezu sa. Preto ich odstranim a nahradim
    # TECUCOU 2-stlpcovou tabulkou bez okrajov (w:cantSplit) podla FIA standardu.
    body = d.element.body
    for r in list(body.iter(qn('w:r'))):
        if any(ch.tag.split('}')[-1] in ('drawing', 'pict', 'AlternateContent', 'object')
               for ch in r.iter()):
            par = r.getparent()
            if par is not None:
                par.remove(r)

    # 4) banner vpravo hore — PLAVAJUCI (ukotveny), nezabera riadok v texte
    bn = os.path.join(HERE, "banner_%s.png" % code)
    tgt = d.paragraphs[1] if not d.paragraphs[1].text.strip() else d.paragraphs[0]
    pf = tgt.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None
    run = tgt.add_run()
    run.add_picture(bn, width=Cm(3.6))
    _float_right(run)

    # 5) podpisovy blok — 2-stlpcova tabulka BEZ okrajov, drzi pokope, teie s textom
    tbl = d.add_table(rows=1, cols=2)
    tbl.autofit = False
    tblPr = tbl._element.find(qn('w:tblPr'))
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'none'); e.set(qn('w:sz'), '0')
        borders.append(e)
    tblPr.append(borders)
    trPr = OxmlElement('w:trPr'); trPr.append(OxmlElement('w:cantSplit'))
    tbl.rows[0]._element.insert(0, trPr)
    L, R = tbl.rows[0].cells
    L.width = Cm(7.4); R.width = Cm(9.2)
    for c in (L, R):
        c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    L.paragraphs[0].add_run(BOXES.get("S úctou", "S úctou"))
    L.add_paragraph()
    L.add_paragraph().add_run(BOXES.get("V Kumhausene 24.12.2025", "Kumhausen, 24.12.2025"))
    ps = R.paragraphs[0]
    ps.paragraph_format.left_indent = Cm(1.1)
    ps.paragraph_format.space_after = Pt(0)
    ps.add_run().add_picture(os.path.join(HERE, "sig_hires_clean.png"), width=Cm(3.4))
    rn = R.add_paragraph().add_run("Peter Ferenc ...............................")
    rn.bold = True

    d.save(out_docx)

    # 6) FIT-PASS — FR/ES/IT su dlhsie ako SK, original ma 2 strany → klon musi mat tiez 2.
    # Stlacame LEN prazdne medzery, odstupy a riadkovanie; text ostava.
    def para_mark_size(p, pt):
        """velkost ZNACKY odseku (prazdny odsek nema behy, vysku urcuje w:pPr/w:rPr/w:sz)"""
        pPr = p._element.get_or_add_pPr()
        rPr = pPr.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr'); pPr.insert(0, rPr)
        for tag in ('w:sz', 'w:szCs'):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag); rPr.append(el)
            el.set(qn('w:val'), str(int(pt * 2)))

    LEVELS = [None,
              (8, None, None), (6, 6, None), (4, 4, None),
              (3, 2, 1.0), (2, 0, 1.0), (2, 0, 0.96)]
    for lvl in LEVELS:
        if lvl is not None:
            spacer_pt, gap_pt, ls = lvl
            d = Document(out_docx)
            for p in d.paragraphs:
                if p._element.findall('.//' + qn('w:drawing')):
                    continue
                if not p.text.strip():
                    para_mark_size(p, spacer_pt)
                    for r in p.runs:
                        r.font.size = Pt(spacer_pt)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                else:
                    if gap_pt is not None:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(gap_pt)
                    if ls is not None:
                        p.paragraph_format.line_spacing = ls
            d.save(out_docx)
        subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", HERE, out_docx],
                       check=True, capture_output=True, timeout=240)
        n = len(PdfReader(out_docx.replace(".docx", ".pdf")).pages)
        print("   fit %s → %d strán" % (lvl, n))
        if n <= 2:
            break
    print("hotovo:", out_docx.replace(".docx", ".pdf"), "| strán:", n)

if __name__ == "__main__":
    build(sys.argv[1])
